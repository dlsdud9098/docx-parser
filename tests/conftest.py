"""Pytest configuration and fixtures for docx-parser tests.

This module provides shared fixtures, mocks, and configuration for the
test suite. Fixtures are organized by scope and purpose.

Fixture Scopes:
    - session: Created once per test session (e.g., golden files)
    - function: Created fresh for each test (default, e.g., mocks)

Mock Strategy:
    - External APIs (OpenAI, Anthropic, etc.): Always mock
    - External processes (subprocess): Always mock
    - Standard library (zipfile, PIL): Prefer real implementations
    - File system: Use tmp_path fixture
"""

from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any, Dict, Generator, List, Optional
from unittest.mock import MagicMock, patch

import pytest


# =============================================================================
# Path Fixtures
# =============================================================================

@pytest.fixture(scope="session")
def fixtures_dir() -> Path:
    """Return the path to the fixtures directory."""
    return Path(__file__).parent / "fixtures"


@pytest.fixture(scope="session")
def golden_dir(fixtures_dir: Path) -> Path:
    """Return the path to the golden files directory."""
    return fixtures_dir / "golden"


# =============================================================================
# DOCX Generation Fixtures
# =============================================================================

@pytest.fixture(scope="function")
def sample_docx(tmp_path: Path) -> Path:
    """Create a simple DOCX file for testing.

    This fixture creates a minimal valid DOCX with basic content
    using python-docx.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("Content in section 1.")

        path = tmp_path / "test_document.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture(scope="function")
def sample_docx_with_images(tmp_path: Path) -> Path:
    """Create a DOCX file with embedded images.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image
        import io

        # Create a simple test image
        img = Image.new("RGB", (100, 100), color="red")
        img_bytes = io.BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        # Save image to temp file for adding to document
        img_path = tmp_path / "test_image.png"
        img.save(str(img_path))

        doc = Document()
        doc.add_heading("Document with Images", level=1)
        doc.add_paragraph("Here is an image:")
        doc.add_picture(str(img_path), width=Inches(1.0))
        doc.add_paragraph("After the image.")

        path = tmp_path / "test_with_images.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx or Pillow not installed")


@pytest.fixture(scope="function")
def sample_docx_with_tables(tmp_path: Path) -> Path:
    """Create a DOCX file with tables.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Document with Tables", level=1)

        # Add a simple table
        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        # Fill the table
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Row {i+1}, Col {j+1}"

        doc.add_paragraph("After the table.")

        path = tmp_path / "test_with_tables.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture(scope="function")
def sample_docx_with_headings(tmp_path: Path) -> Path:
    """Create a DOCX file with multiple heading levels.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Title (Level 1)", level=1)
        doc.add_paragraph("Introduction paragraph.")

        doc.add_heading("Chapter 1 (Level 2)", level=2)
        doc.add_paragraph("Chapter 1 content.")

        doc.add_heading("Section 1.1 (Level 3)", level=3)
        doc.add_paragraph("Section 1.1 content.")

        doc.add_heading("Chapter 2 (Level 2)", level=2)
        doc.add_paragraph("Chapter 2 content.")

        path = tmp_path / "test_with_headings.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture(scope="function")
def empty_docx(tmp_path: Path) -> Path:
    """Create an empty DOCX file.

    Returns:
        Path to the created DOCX file.
    """
    try:
        from docx import Document

        doc = Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture(scope="function")
def corrupted_docx(tmp_path: Path) -> Path:
    """Create a corrupted DOCX file (invalid ZIP).

    Returns:
        Path to the corrupted file.
    """
    path = tmp_path / "corrupted.docx"
    path.write_text("This is not a valid DOCX file")
    return path


@pytest.fixture(scope="function")
def non_docx_file(tmp_path: Path) -> Path:
    """Create a text file with .docx extension.

    Returns:
        Path to the invalid file.
    """
    path = tmp_path / "fake.docx"
    path.write_text("Plain text content, not a DOCX")
    return path


# =============================================================================
# Mock Fixtures - Vision Providers
# =============================================================================

@pytest.fixture(scope="function")
def mock_openai_client() -> Generator[MagicMock, None, None]:
    """Mock the OpenAI client for vision tests.

    Yields:
        Mocked OpenAI client instance.
    """
    with patch("openai.OpenAI") as mock_class:
        mock_client = MagicMock()

        # Mock chat completions response
        mock_response = MagicMock()
        mock_response.choices = [MagicMock()]
        mock_response.choices[0].message.content = "A test image description"
        mock_response.usage.total_tokens = 100

        mock_client.chat.completions.create.return_value = mock_response
        mock_class.return_value = mock_client

        yield mock_client


@pytest.fixture(scope="function")
def mock_anthropic_client() -> Generator[MagicMock, None, None]:
    """Mock the Anthropic client for vision tests.

    Yields:
        Mocked Anthropic client instance.
    """
    with patch("anthropic.Anthropic") as mock_class:
        mock_client = MagicMock()

        # Mock messages response
        mock_response = MagicMock()
        mock_response.content = [MagicMock()]
        mock_response.content[0].text = "A test image description"
        mock_response.usage.output_tokens = 50

        mock_client.messages.create.return_value = mock_response
        mock_class.return_value = mock_client

        yield mock_client


@pytest.fixture(scope="function")
def mock_google_genai() -> Generator[MagicMock, None, None]:
    """Mock the Google Generative AI client for vision tests.

    Yields:
        Mocked GenerativeModel instance.
    """
    with patch("google.generativeai.GenerativeModel") as mock_class:
        mock_model = MagicMock()

        # Mock generate_content response
        mock_response = MagicMock()
        mock_response.text = "A test image description"

        mock_model.generate_content.return_value = mock_response
        mock_class.return_value = mock_model

        yield mock_model


# =============================================================================
# Mock Fixtures - External Processes
# =============================================================================

@pytest.fixture(scope="function")
def mock_jxrdecapp() -> Generator[MagicMock, None, None]:
    """Mock the JxrDecApp subprocess for WDP conversion tests.

    Yields:
        Tuple of (mock_which, mock_run) for subprocess operations.
    """
    with patch("shutil.which") as mock_which:
        with patch("subprocess.run") as mock_run:
            mock_which.return_value = "/usr/bin/JxrDecApp"

            mock_result = MagicMock()
            mock_result.returncode = 0
            mock_run.return_value = mock_result

            yield mock_run


@pytest.fixture(scope="function")
def mock_jxrdecapp_not_installed() -> Generator[MagicMock, None, None]:
    """Mock JxrDecApp as not installed.

    Yields:
        Mocked shutil.which returning None.
    """
    with patch("shutil.which") as mock_which:
        mock_which.return_value = None
        yield mock_which


# =============================================================================
# Mock Fixtures - PIL
# =============================================================================

@pytest.fixture(scope="function")
def mock_pil_image() -> Generator[MagicMock, None, None]:
    """Mock PIL Image for image processing tests.

    Yields:
        Mocked PIL.Image.open function.
    """
    with patch("PIL.Image.open") as mock_open:
        mock_img = MagicMock()
        mock_img.mode = "RGB"
        mock_img.size = (100, 100)
        mock_img.format = "PNG"

        mock_open.return_value = mock_img
        yield mock_open


# =============================================================================
# Helper Fixtures
# =============================================================================

@pytest.fixture(scope="function")
def output_dir(tmp_path: Path) -> Path:
    """Create a temporary output directory.

    Returns:
        Path to the output directory.
    """
    output = tmp_path / "output"
    output.mkdir()
    return output


@pytest.fixture(scope="function")
def sample_image_bytes() -> bytes:
    """Create sample PNG image bytes for testing.

    Returns:
        Raw PNG image bytes.
    """
    try:
        from PIL import Image
        import io

        img = Image.new("RGB", (10, 10), color="blue")
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        return buffer.getvalue()
    except ImportError:
        # Return minimal valid PNG bytes
        return (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00"
            b"\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18"
            b"\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
        )


@pytest.fixture(scope="function")
def sample_jpeg_bytes() -> bytes:
    """Create sample JPEG image bytes for testing.

    Returns:
        Raw JPEG image bytes.
    """
    try:
        from PIL import Image
        import io

        img = Image.new("RGB", (10, 10), color="green")
        buffer = io.BytesIO()
        img.save(buffer, format="JPEG")
        return buffer.getvalue()
    except ImportError:
        # Return minimal JPEG header
        return b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00"


# =============================================================================
# Marker Registration
# =============================================================================

def pytest_configure(config: pytest.Config) -> None:
    """Register custom markers."""
    config.addinivalue_line("markers", "slow: marks tests as slow")
    config.addinivalue_line("markers", "requires_api: tests requiring external API")
    config.addinivalue_line("markers", "integration: integration tests")
    config.addinivalue_line("markers", "e2e: end-to-end tests")
