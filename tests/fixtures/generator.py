"""DOCX fixture generator for tests.

This module provides functions to generate various DOCX test fixtures
programmatically using python-docx. These fixtures can be used for
unit tests, integration tests, and edge case testing.

Usage:
    from tests.fixtures.generator import create_simple_docx
    create_simple_docx(Path("test.docx"))
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


def _require_docx() -> None:
    """Raise ImportError if python-docx is not available."""
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is required to generate DOCX fixtures. "
            "Install with: pip install python-docx"
        )


def _require_pil() -> None:
    """Raise ImportError if Pillow is not available."""
    if not HAS_PIL:
        raise ImportError(
            "Pillow is required to generate image fixtures. "
            "Install with: pip install Pillow"
        )


# =============================================================================
# Basic Document Generators
# =============================================================================

def create_simple_docx(path: Path) -> Path:
    """Create a simple DOCX with basic text content.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Simple Test Document", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.save(str(path))
    return path


def create_empty_docx(path: Path) -> Path:
    """Create an empty DOCX file.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.save(str(path))
    return path


def create_docx_with_metadata(
    path: Path,
    title: str = "Test Title",
    author: str = "Test Author",
    subject: str = "Test Subject",
    keywords: str = "test, docx, parser",
) -> Path:
    """Create a DOCX with custom metadata.

    Args:
        path: Output path for the DOCX file.
        title: Document title.
        author: Document author.
        subject: Document subject.
        keywords: Document keywords.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.subject = subject
    doc.core_properties.keywords = keywords

    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Author: {author}")
    doc.save(str(path))
    return path


# =============================================================================
# Heading Document Generators
# =============================================================================

def create_docx_with_headings(
    path: Path,
    max_level: int = 6,
) -> Path:
    """Create a DOCX with multiple heading levels.

    Args:
        path: Output path for the DOCX file.
        max_level: Maximum heading level to include (1-6).

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()

    for level in range(1, min(max_level + 1, 7)):
        doc.add_heading(f"Heading Level {level}", level=level)
        doc.add_paragraph(f"Content under heading level {level}.")

    doc.save(str(path))
    return path


def create_docx_with_font_size_headings(
    path: Path,
    sizes: Optional[List[Tuple[int, str]]] = None,
) -> Path:
    """Create a DOCX with headings defined by font sizes (no styles).

    Args:
        path: Output path for the DOCX file.
        sizes: List of (font_size_pt, text) tuples.

    Returns:
        The path to the created file.
    """
    _require_docx()

    if sizes is None:
        sizes = [
            (24, "Large Title (24pt)"),
            (18, "Section Header (18pt)"),
            (14, "Subsection (14pt)"),
            (12, "Normal text paragraph"),
            (16, "Another Header (16pt)"),
            (12, "More normal text"),
        ]

    doc = Document()

    for font_size, text in sizes:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(font_size)

    doc.save(str(path))
    return path


# =============================================================================
# Table Document Generators
# =============================================================================

def create_docx_with_simple_table(
    path: Path,
    rows: int = 3,
    cols: int = 3,
) -> Path:
    """Create a DOCX with a simple table.

    Args:
        path: Output path for the DOCX file.
        rows: Number of table rows.
        cols: Number of table columns.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Document with Table", level=1)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # Header row
    for j, cell in enumerate(table.rows[0].cells):
        cell.text = f"Header {j + 1}"

    # Data rows
    for i in range(1, rows):
        for j, cell in enumerate(table.rows[i].cells):
            cell.text = f"R{i}C{j + 1}"

    doc.save(str(path))
    return path


def create_docx_with_merged_cells(path: Path) -> Path:
    """Create a DOCX with vertically and horizontally merged cells.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Table with Merged Cells", level=1)

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    # Fill all cells first
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i+1}C{j+1}"

    # Horizontal merge: merge cells in first row
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Merged H"

    # Vertical merge: merge cells in first column
    table.cell(1, 0).merge(table.cell(2, 0))
    table.cell(1, 0).text = "Merged V"

    doc.save(str(path))
    return path


def create_docx_with_nested_tables(path: Path) -> Path:
    """Create a DOCX with nested tables.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Nested Tables", level=1)

    # Outer table
    outer_table = doc.add_table(rows=2, cols=2)
    outer_table.style = "Table Grid"

    outer_table.cell(0, 0).text = "Cell 1"
    outer_table.cell(0, 1).text = "Cell 2"
    outer_table.cell(1, 1).text = "Cell 4"

    # Add inner table to cell (1, 0)
    inner_cell = outer_table.cell(1, 0)
    inner_cell.text = ""  # Clear default text
    inner_table = inner_cell.add_table(rows=2, cols=2)

    for i, row in enumerate(inner_table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Inner {i+1},{j+1}"

    doc.save(str(path))
    return path


# =============================================================================
# Image Document Generators
# =============================================================================

def create_docx_with_images(
    path: Path,
    num_images: int = 2,
    image_size: Tuple[int, int] = (100, 100),
) -> Path:
    """Create a DOCX with embedded images.

    Args:
        path: Output path for the DOCX file.
        num_images: Number of images to embed.
        image_size: Size of generated images (width, height).

    Returns:
        The path to the created file.
    """
    _require_docx()
    _require_pil()

    doc = Document()
    doc.add_heading("Document with Images", level=1)

    colors = ["red", "green", "blue", "yellow", "purple"]

    for i in range(num_images):
        doc.add_paragraph(f"Image {i + 1}:")

        # Create a colored image
        color = colors[i % len(colors)]
        img = Image.new("RGB", image_size, color=color)

        # Save to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format="PNG")
        img_bytes.seek(0)

        # Add to document
        doc.add_picture(img_bytes, width=Inches(1.5))

    doc.save(str(path))
    return path


# =============================================================================
# Unicode and Special Content Generators
# =============================================================================

def create_docx_with_unicode(path: Path) -> Path:
    """Create a DOCX with various Unicode content.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Unicode Test Document", level=1)

    # Korean
    doc.add_heading("한글 텍스트", level=2)
    doc.add_paragraph("안녕하세요. 이것은 한글 테스트입니다.")

    # Chinese
    doc.add_heading("中文文本", level=2)
    doc.add_paragraph("你好。这是中文测试。")

    # Japanese
    doc.add_heading("日本語テキスト", level=2)
    doc.add_paragraph("こんにちは。これは日本語のテストです。")

    # Emoji
    doc.add_heading("Emoji Test", level=2)
    doc.add_paragraph("Hello! 😀 👍 🎉 🚀 ❤️")

    # Math symbols
    doc.add_heading("Math Symbols", level=2)
    doc.add_paragraph("α β γ δ ∑ ∏ ∫ √ ∞ ≠ ≤ ≥")

    doc.save(str(path))
    return path


def create_docx_with_special_chars(path: Path) -> Path:
    """Create a DOCX with special characters that may cause issues.

    Args:
        path: Output path for the DOCX file.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Special Characters Test", level=1)

    # Characters that may need escaping
    doc.add_paragraph("Pipe: | Backslash: \\ Asterisk: * Underscore: _")
    doc.add_paragraph("Brackets: [test] {test} <test>")
    doc.add_paragraph("Quotes: 'single' \"double\" `backtick`")
    doc.add_paragraph("Hash: # Dollar: $ Percent: % Caret: ^")
    doc.add_paragraph("Ampersand: & At: @ Tilde: ~")

    # Newlines and tabs in text
    doc.add_paragraph("Line with\ttab character")
    doc.add_paragraph("Line 1\nLine 2 (embedded newline)")

    doc.save(str(path))
    return path


# =============================================================================
# Large Document Generators
# =============================================================================

def create_large_docx(
    path: Path,
    num_paragraphs: int = 1000,
    words_per_paragraph: int = 100,
) -> Path:
    """Create a large DOCX for performance testing.

    Args:
        path: Output path for the DOCX file.
        num_paragraphs: Number of paragraphs to generate.
        words_per_paragraph: Words per paragraph.

    Returns:
        The path to the created file.
    """
    _require_docx()

    doc = Document()
    doc.add_heading("Large Test Document", level=1)

    sample_words = [
        "lorem", "ipsum", "dolor", "sit", "amet", "consectetur",
        "adipiscing", "elit", "sed", "do", "eiusmod", "tempor"
    ]

    for i in range(num_paragraphs):
        if i % 100 == 0:
            doc.add_heading(f"Section {i // 100 + 1}", level=2)

        # Generate paragraph content
        words = []
        for j in range(words_per_paragraph):
            words.append(sample_words[j % len(sample_words)])
        doc.add_paragraph(" ".join(words))

    doc.save(str(path))
    return path


# =============================================================================
# Corrupted/Invalid Document Generators
# =============================================================================

def create_corrupted_zip(path: Path) -> Path:
    """Create a file that looks like a DOCX but has invalid content.

    Args:
        path: Output path for the corrupted file.

    Returns:
        The path to the created file.
    """
    # Write random bytes that are not a valid ZIP
    path.write_bytes(b"PK\x03\x04" + b"\x00" * 100)  # Invalid ZIP
    return path


def create_missing_document_xml(path: Path) -> Path:
    """Create a DOCX-like ZIP without document.xml.

    Args:
        path: Output path for the invalid DOCX.

    Returns:
        The path to the created file.
    """
    import zipfile

    with zipfile.ZipFile(path, "w") as zf:
        # Add content types (required)
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>'
        )
        # Missing word/document.xml

    return path


# =============================================================================
# All-in-one Generator
# =============================================================================

def generate_all_fixtures(output_dir: Path) -> None:
    """Generate all test fixtures in the specified directory.

    Args:
        output_dir: Directory to save all fixtures.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    fixtures = [
        ("simple.docx", create_simple_docx),
        ("empty.docx", create_empty_docx),
        ("with_headings.docx", create_docx_with_headings),
        ("with_simple_table.docx", create_docx_with_simple_table),
        ("with_merged_cells.docx", create_docx_with_merged_cells),
        ("with_unicode.docx", create_docx_with_unicode),
        ("with_special_chars.docx", create_docx_with_special_chars),
    ]

    # Add fixtures requiring PIL
    if HAS_PIL:
        fixtures.append(("with_images.docx", create_docx_with_images))

    for filename, generator in fixtures:
        try:
            generator(output_dir / filename)
            print(f"Created: {filename}")
        except Exception as e:
            print(f"Failed to create {filename}: {e}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        output = Path(sys.argv[1])
    else:
        output = Path("tests/fixtures/generated")

    generate_all_fixtures(output)
